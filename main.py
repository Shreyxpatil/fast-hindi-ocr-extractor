"""
Fast Hindi + English Text Extraction Server

Two-layer strategy for all file types:
  1. Digital text extraction (instant) — for PDFs with embedded text, DOCX, TXT
  2. Tesseract OCR fallback (fast) — for scanned PDFs, images, KrutiDev-encoded pages

Supports: PDF, DOC, DOCX, TXT, PNG, JPG, JPEG
"""

import os
import io
import re
import time
import json
import tempfile
import asyncio
import traceback
import base64
import uuid
from typing import List, Dict, Tuple
from concurrent.futures import ThreadPoolExecutor

os.environ["OMP_THREAD_LIMIT"] = "1"  # Prevent Tesseract thread contention

import uvicorn
from fastapi import FastAPI, UploadFile, File, Form
import fitz  # PyMuPDF
from PIL import Image
import pytesseract
from docx import Document

app = FastAPI()

OUTPUT_DIR = "outputs"
os.makedirs(OUTPUT_DIR, exist_ok=True)
os.makedirs(os.path.join(OUTPUT_DIR, "images"), exist_ok=True)
os.makedirs(os.path.join(OUTPUT_DIR, "tables"), exist_ok=True)

# Thread pool for parallel Tesseract OCR
OCR_WORKERS = min(os.cpu_count() or 4, 8)


# ============================================================
# DETECTION UTILS
# ============================================================

KRUTIDEV_MARKERS = set("÷Ê⁄Ã∑§ﬂ◊ÿ¥Á‹ŸÒœ»›'ﬁ¬˝·Ë™Ù¢Œ†‡")



# ============================================================
# HELPER FUNCTIONS
# ============================================================

def save_base64_to_file(data: str, folder: str, file_id: str) -> str:
    """Save Base64 data to a text file and return the relative path."""
    filename = f"{file_id}.txt"
    path = os.path.join(OUTPUT_DIR, folder, filename)
    with open(path, "w", encoding="utf-8") as f:
        f.write(data)
    return path

def _collapse_devanagari_repeats(text: str) -> str:
    """Collapse repeating Devanagari syllable patterns like 'निनिनिनिनि' -> 'नि'."""
    if not text: return text
    # Simple heuristic: if a substring of length > 2 repeats > 3 times consecutively
    # Actually, a safer approach is to remove lines that are purely repetitive garbage
    # or use a regex for specific patterns.
    # For now, let's target the known 'ni-ni-ni' pattern
    return re.sub(r'(..+?)\1{4,}', r'\1', text)

def _is_garbage_token(token: str) -> bool:
    """Check if a token looks like OCR garbage."""
    if len(token) > 30: return True # Super long tokens are usually garbage
    # High consonant density check for Latin could go here
    # Repeating characters
    if re.search(r'(.)\1{4,}', token): return True
    return False

def clean_ocr_text(text: str) -> str:
    """Clean OCR output by removing garbage and normalizing whitespace."""
    if not text: return ""
    
    # normalize whitespace
    text = re.sub(r'\s+', ' ', text).strip()
    
    # specific garbage patterns
    text = re.sub(r'[cceess]{5,}', '', text, flags=re.IGNORECASE) # 'ceccseses...'
    text = _collapse_devanagari_repeats(text)
    
    tokens = text.split()
    clean_tokens = [t for t in tokens if not _is_garbage_token(t)]
    return " ".join(clean_tokens)

def count_hindi(text: str) -> int:
    """Count Devanagari characters in text."""
    return sum(1 for c in text if '\u0900' <= c <= '\u097f')

def has_content(text: str) -> bool:
    """Check if text has meaningful content."""
    return len(text.strip()) > 0

def is_krutidev(text: str) -> bool:
    """Check if text looks like KrutiDev encoding (garbage to UTF-8 eyes)."""
    if not text: return False
    # Heuristic: if text has high overlap with KrutiDev markers and low valid Devanagari
    chars = set(text)
    overlap = len(chars.intersection(KRUTIDEV_MARKERS))
    has_hindi = count_hindi(text) > 5
    return overlap > 2 and not has_hindi

ENGLISH_STOP_WORDS = {
    "the", "and", "of", "to", "in", "is", "that", "it", "as", "was", "for",
    "on", "are", "with", "be", "this", "from", "at", "or", "by", "an", "not",
    "what", "all", "were", "we", "when", "your", "can", "said", "there", "use",
    "date", "page", "no", "yes", "if"
}

def is_bad_text_layer(text: str) -> bool:
    """
    Check if text layer is likely garbage (e.g. corrupted ASCII in Hindi doc).
    True = Bad (Force OCR), False = Good (Keep digital text).
    """
    if not text: return True
    
    # If it has meaningful Hindi, it's good
    if count_hindi(text) > 5:
        return False
        
    # If it has meaningful English words, it's good
    words = text.lower().split()
    if not words: return True
    
    # Calculate stop word ratio
    stop_count = sum(1 for w in words if w in ENGLISH_STOP_WORDS)
    ratio = stop_count / len(words)
    
    # If text is long (>50 words) but has very low English stop word density (<10%), 
    # it's likely corrupted or garbage text (e.g., Page 1 of mixed doc).
    if len(words) > 50 and ratio < 0.10:
        return True

    # If text is long (>50 chars) but has NO stop words, it's likely garbage
    # Example: "122 bOp-S) IbbS bible..."
    if len(text) > 50 and stop_count == 0:
        return True
        
    return False

def tesseract_ocr_page(image: Image.Image) -> str:
    """Run Tesseract on a single image."""
    try:
        # Hindi + English
        return pytesseract.image_to_string(image, lang="hin+eng", config="--psm 1")
    except Exception as e:
        print(f"  ⚠️ OCR Error: {e}")
        return ""

def tesseract_ocr_batch(images: List[Image.Image]) -> List[str]:
    """Run Tesseract on multiple images in parallel."""
    with ThreadPoolExecutor(max_workers=OCR_WORKERS) as executor:
        return list(executor.map(tesseract_ocr_page, images))


def process_pdf_sync(contents: bytes, file_name: str, document_id: str) -> Tuple[List[str], List[Dict], List[Dict], List[Dict]]:
    """
    Process a PDF file.
    Phase 1: Extract digital text from each page (instant).
    Phase 2: OCR pages that have no text or garbled KrutiDev text.
    Phase 3: Extract tables and images as Base64.
    """
    start = time.time()

    with fitz.open(stream=contents, filetype="pdf") as doc:
        total_pages = len(doc)
        print(f"  📄 PDF: {total_pages} pages")

        texts = [""] * total_pages
        clean_count = 0
        needs_ocr = []
        
        extracted_tables = []
        extracted_images = []

        # --- Phase 1: Digital text extraction (instant) ---
        t1 = time.time()
        for i in range(total_pages):
            page = doc[i]
            digital = page.get_text()

            # --- Table Extraction ---
            try:
                tables = page.find_tables()
                for tab in tables:
                    # Clip coordinates
                    bbox = tab.bbox
                    # Render table area to image
                    pix = page.get_pixmap(clip=bbox, dpi=150)
                    img_data = pix.tobytes("png")
                    b64_str = base64.b64encode(img_data).decode("utf-8")
                    
                    table_id = f"TBL_{uuid.uuid4().hex[:8]}"
                    file_path = save_base64_to_file(b64_str, "tables", table_id)
                    
                    extracted_tables.append({
                        "table_id": table_id,
                        "file_name": file_name,
                        "page_number": i + 1,
                        "text": "\\n".join(["\\t".join(cell) for cell in tab.extract() if cell]), # Simple text rep
                        "table_image_path": file_path,
                    })
            except Exception as e:
                print(f"  ⚠️ Table extraction error on page {i+1}: {e}")

            # --- Image Extraction (Scanning for embedded images) ---
            try:
                img_list = page.get_images(full=True)
                for img_index, img in enumerate(img_list):
                    xref = img[0]
                    base_image = doc.extract_image(xref)
                    image_bytes = base_image["image"]
                    
                    if len(image_bytes) < 5000: # Skip tiny icons/lines
                        continue
                        
                    b64_str = base64.b64encode(image_bytes).decode("utf-8")
                    img_id = f"IMG_{uuid.uuid4().hex[:8]}"
                    file_path = save_base64_to_file(b64_str, "images", img_id)
                    
                    extracted_images.append({
                        "image_id": img_id,
                        "file_name": file_name,
                        "page_number": i + 1,
                        "image_path": file_path
                    })
            except Exception as e:
                 print(f"  ⚠️ Image extraction error on page {i+1}: {e}")


            if not has_content(digital):
                needs_ocr.append(i)
            elif is_krutidev(digital):
                needs_ocr.append(i)
            elif is_bad_text_layer(digital):
                needs_ocr.append(i)
            elif count_hindi(digital) > 0 or len(digital.strip()) > 50:
                texts[i] = digital
                clean_count += 1
            else:
                needs_ocr.append(i)

        print(f"  ⚡ Phase 1 ({time.time()-t1:.1f}s): {clean_count} clean digital pages")
        print(f"  🔍 Phase 2: {len(needs_ocr)} pages need Tesseract OCR")
        print(f"  📊 Phase 3: Extracted {len(extracted_tables)} tables, {len(extracted_images)} images")


        # --- Phase 2: Tesseract OCR for remaining pages ---
        if needs_ocr:
            t2 = time.time()
            images = []
            for idx in needs_ocr:
                pix = doc[idx].get_pixmap(dpi=150)
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                images.append(img)
            print(f"  📸 Rendered {len(images)} pages in {time.time()-t2:.1f}s")

            t3 = time.time()
            ocr_results = tesseract_ocr_batch(images)
            for j, idx in enumerate(needs_ocr):
                texts[idx] = ocr_results[j]

            ok = sum(1 for r in ocr_results if r.strip())
            elapsed_ocr = max(time.time() - t3, 0.1)
            print(f"  ⏱️ Phase 2: {ok}/{len(needs_ocr)} pages in {elapsed_ocr:.1f}s ({len(needs_ocr)/elapsed_ocr:.1f} p/s)")

        all_metadata = [
            {"file_name": file_name, "document_id": document_id, "page_num": i + 1, "extension": "pdf"}
            for i in range(total_pages)
        ]

        # Clean OCR garbage from all pages
        texts = [clean_ocr_text(t) for t in texts]

        total_hindi = sum(count_hindi(t) for t in texts)
        print(f"  ✅ PDF done: {total_pages}p | {total_hindi:,} Hindi | {time.time()-start:.1f}s")
        return texts, all_metadata, extracted_tables, extracted_images

async def process_pdf(contents: bytes, file_name: str, document_id: str) -> Tuple[List[str], List[Dict], List[Dict], List[Dict]]:
    """Async wrapper — runs PDF processing in a separate thread."""
    return await asyncio.to_thread(process_pdf_sync, contents, file_name, document_id)


# ============================================================
# DOC PROCESSOR (.doc — legacy Microsoft Word)
# ============================================================

def process_doc_sync(contents: bytes, file_name: str, document_id: str) -> Tuple[List[str], List[Dict], List[Dict], List[Dict]]:
    """
    Process a .doc file.
    Phase 1: Try text extraction via antiword or textract.
    Phase 2: If garbled or empty, convert to images and OCR.
    """
    start = time.time()
    print(f"  📄 DOC: {file_name}")

    # Save to temp file
    with tempfile.NamedTemporaryFile(delete=False, suffix=".doc") as tmp:
        tmp.write(contents)
        tmp_path = tmp.name

    try:
        # Phase 1: Try extracting text directly
        # .doc is binary — try reading as raw text (may work for simple docs)
        text = ""
        try:
            raw = contents.decode("utf-8", errors="ignore")
            # Filter out binary garbage — keep only printable chars
            text = re.sub(r'[^\x20-\x7E\u0900-\u097F\n\r\t]', '', raw)
            text = re.sub(r'\n{3,}', '\n\n', text).strip()
        except Exception:
            pass

        if has_content(text) and not is_krutidev(text):
            print(f"  ⚡ Phase 1: extracted {len(text)} chars directly")
        else:
            # Phase 2: Convert .doc to PDF via LibreOffice, then process as PDF
            print(f"  🔍 Phase 2: converting .doc → PDF for OCR")
            try:
                import subprocess
                result = subprocess.run(
                    ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", tempfile.gettempdir(), tmp_path],
                    capture_output=True, timeout=30
                )
                pdf_path = tmp_path.replace(".doc", ".pdf")
                if os.path.exists(pdf_path):
                    with open(pdf_path, "rb") as f:
                        pdf_bytes = f.read()
                    texts, metadata, tables, images = process_pdf_sync(pdf_bytes, file_name, document_id)
                    # Fix extension in metadata
                    for m in metadata:
                        m["extension"] = "doc"
                    os.unlink(pdf_path)
                    os.unlink(tmp_path)
                    return texts, metadata, tables, images
            except Exception as e:
                print(f"  ⚠️ LibreOffice conversion failed: {e}")
                # Fallback: just return whatever text we got
                pass

        text = clean_ocr_text(text)
        total_hindi = count_hindi(text)
        print(f"  ✅ DOC done: {total_hindi:,} Hindi | {time.time()-start:.1f}s")
        return [text], [{"file_name": file_name, "document_id": document_id, "page_num": 1, "extension": "doc"}], [], []
    finally:
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)


async def process_doc(contents: bytes, file_name: str, document_id: str) -> Tuple[List[str], List[Dict], List[Dict], List[Dict]]:
    """Async wrapper — runs DOC processing in a separate thread."""
    return await asyncio.to_thread(process_doc_sync, contents, file_name, document_id)


# ============================================================
# DOCX PROCESSOR (.docx — modern Microsoft Word)
# ============================================================

def process_docx_sync(contents: bytes, file_name: str, document_id: str) -> Tuple[List[str], List[Dict], List[Dict], List[Dict]]:
    """
    Process a .docx file.
    Phase 1: Extract text from paragraphs and tables.
    Phase 2: If garbled/KrutiDev, render pages and OCR.
    """
    start = time.time()
    print(f"  📄 DOCX: {file_name}")

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp.write(contents)
        tmp_path = tmp.name

    try:
        doc = Document(tmp_path)

        # Phase 1: Extract text from paragraphs + tables
        parts = []

        # Paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)

        # Tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    parts.append(row_text)

        text = "\n".join(parts)
        print(f"  ⚡ Phase 1: {len(parts)} paragraphs/rows, {len(text)} chars")

        # Check if text is garbled (KrutiDev)
        if has_content(text) and is_krutidev(text):
            print(f"  🔍 Phase 2: KrutiDev detected — converting DOCX → PDF for OCR")
            try:
                import subprocess
                result = subprocess.run(
                    ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", tempfile.gettempdir(), tmp_path],
                    capture_output=True, timeout=60
                )
                pdf_path = tmp_path.replace(".docx", ".pdf")
                if os.path.exists(pdf_path):
                    with open(pdf_path, "rb") as f:
                        pdf_bytes = f.read()
                    texts, metadata, tables, images = process_pdf_sync(pdf_bytes, file_name, document_id)
                    for m in metadata:
                        m["extension"] = "docx"
                    os.unlink(pdf_path)
                    os.unlink(tmp_path)
                    return texts, metadata, tables, images
            except Exception as e:
                print(f"  ⚠️ LibreOffice conversion failed: {e}, using raw text")

        text = clean_ocr_text(text)
        total_hindi = count_hindi(text)
        print(f"  ✅ DOCX done: {total_hindi:,} Hindi | {time.time()-start:.1f}s")
        return [text], [{"file_name": file_name, "document_id": document_id, "page_num": 1, "extension": "docx"}], [], []
    finally:
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)


async def process_docx(contents: bytes, file_name: str, document_id: str) -> Tuple[List[str], List[Dict], List[Dict], List[Dict]]:
    """Async wrapper — runs DOCX processing in a separate thread."""
    return await asyncio.to_thread(process_docx_sync, contents, file_name, document_id)


# ============================================================
# TXT PROCESSOR (.txt — plain text)
# ============================================================

def process_txt_sync(contents: bytes, file_name: str, document_id: str) -> Tuple[List[str], List[Dict], List[Dict], List[Dict]]:
    """
    Process a .txt file.
    Phase 1: Decode as UTF-8 text.
    Phase 2: If KrutiDev, no OCR possible — return as-is.
    """
    start = time.time()
    print(f"  📄 TXT: {file_name}")

    # Try multiple encodings
    text = ""
    for encoding in ["utf-8", "utf-16", "latin-1"]:
        try:
            text = contents.decode(encoding)
            break
        except (UnicodeDecodeError, Exception):
            continue

    if not text:
        text = contents.decode("utf-8", errors="ignore")

    text = clean_ocr_text(text)
    total_hindi = count_hindi(text)
    print(f"  ✅ TXT done: {len(text):,} chars | {total_hindi:,} Hindi | {time.time()-start:.2f}s")
    return [text], [{"file_name": file_name, "document_id": document_id, "page_num": 1, "extension": "txt"}], [], []


async def process_txt(contents: bytes, file_name: str, document_id: str) -> Tuple[List[str], List[Dict], List[Dict], List[Dict]]:
    """Async wrapper — runs TXT processing in a separate thread."""
    return await asyncio.to_thread(process_txt_sync, contents, file_name, document_id)


# ============================================================
# IMAGE PROCESSOR (.png, .jpg, .jpeg)
# ============================================================

def process_image_sync(contents: bytes, file_name: str, document_id: str, ext: str) -> Tuple[List[str], List[Dict], List[Dict], List[Dict]]:
    """
    Process an image file (PNG, JPG, JPEG).
    Always uses Tesseract OCR — images have no embedded digital text.
    """
    start = time.time()
    print(f"  🖼️ IMAGE: {file_name}")

    image = Image.open(io.BytesIO(contents)).convert("RGB")
    width, height = image.size
    print(f"  📐 Size: {width}x{height}")

    # OCR the image
    text = tesseract_ocr_page(image)
    text = clean_ocr_text(text)
    total_hindi = count_hindi(text)

    print(f"  ✅ IMAGE done: {len(text):,} chars | {total_hindi:,} Hindi | {time.time()-start:.1f}s")
    return [text], [{"file_name": file_name, "document_id": document_id, "page_num": 1, "extension": ext}], [], []


async def process_image(contents: bytes, file_name: str, document_id: str, ext: str) -> Tuple[List[str], List[Dict], List[Dict], List[Dict]]:
    """Async wrapper — runs image OCR in a separate thread."""
    return await asyncio.to_thread(process_image_sync, contents, file_name, document_id, ext)


# ============================================================
# API ROUTE
# ============================================================

SUPPORTED_EXTENSIONS = {
    "pdf": "PDF Document",
    "doc": "Legacy Word Document",
    "docx": "Word Document",
    "txt": "Plain Text",
    "png": "PNG Image",
    "jpg": "JPEG Image",
    "jpeg": "JPEG Image",
}


@app.post("/extract")
async def extract_file(
    file: UploadFile = File(...),
    document_id: str = Form(...)
):
    """Extract text from uploaded file. Supports: PDF, DOC, DOCX, TXT, PNG, JPG, JPEG."""
    try:
        file_name = file.filename or "unknown"
        ext = os.path.splitext(file_name)[1].lower().replace(".", "")
        contents = await file.read()

        start = time.time()
        print(f"\n{'='*60}")
        print(f"📥 {file_name} | {SUPPORTED_EXTENSIONS.get(ext, ext)} | {len(contents)/1024:.0f}KB")
        print(f"{'='*60}")

        if ext not in SUPPORTED_EXTENSIONS:
            return {"error": f"Unsupported file type: .{ext}", "supported": list(SUPPORTED_EXTENSIONS.keys())}

        # Route to the correct processor
        all_text: List[str] = []
        all_metadata: List[Dict] = []
        all_tables: List[Dict] = []
        all_images: List[Dict] = []

        if ext == "pdf":
            all_text, all_metadata, all_tables, all_images = await process_pdf(contents, file_name, document_id)
        elif ext == "doc":
             all_text, all_metadata, all_tables, all_images = await process_doc(contents, file_name, document_id)
        elif ext == "docx":
             all_text, all_metadata, all_tables, all_images = await process_docx(contents, file_name, document_id)
        elif ext == "txt":
             all_text, all_metadata, all_tables, all_images = await process_txt(contents, file_name, document_id)
        elif ext in ("png", "jpg", "jpeg"):
             all_text, all_metadata, all_tables, all_images = await process_image(contents, file_name, document_id, ext)

        # Save output JSON
        output_data = {
            "all_text": all_text,
            "all_metadata": all_metadata,
            "all_tables": all_tables,
            "all_images": all_images
        }

        out_path = os.path.join(OUTPUT_DIR, f"{document_id}.json")
        with open(out_path, "w", encoding="utf-8") as f:
            json.dump(output_data, f, ensure_ascii=False, indent=2)

        elapsed = time.time() - start
        total_hindi = sum(count_hindi(t) for t in all_text)
        total_chars = sum(len(t) for t in all_text)
        print(f"💾 Saved: {out_path} | {len(all_text)}p | {total_chars:,} chars | {total_hindi:,} Hindi | {elapsed:.1f}s")

        return output_data

    except Exception as e:
        print(f"❌ ERROR: {e}")
        traceback.print_exc()
        raise e

if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8000)

